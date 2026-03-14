from __future__ import annotations

import csv
import json
import math
import re
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from textwrap import dedent
from typing import Any, Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(r"c:\Users\Hugo\Desktop\Thesis\Medical-IoMT-")
REPORTS = ROOT / "reports"
OUT_MD = ROOT / "Thesis" / "THESIS_EXPERIMENT_CHAPTER.md"
OUT_DOCX = ROOT / "Thesis" / "THESIS_EXPERIMENT_CHAPTER.docx"


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def read_json(path: Path) -> Any:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def to_float(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        if math.isnan(value) or math.isinf(value):
            return None
        return float(value)
    text = str(value).strip()
    if not text or text.lower() in {"nan", "na", "n/a"}:
        return None
    try:
        number = float(text)
    except ValueError:
        return None
    if math.isnan(number) or math.isinf(number):
        return None
    return number


def clean_text(text: str) -> str:
    lines = [line.strip() for line in dedent(text).strip().splitlines()]
    parts = [line for line in lines if line]
    return " ".join(parts)


def format_count(value: Any) -> str:
    number = to_float(value)
    if number is None:
        return "NA"
    return f"{int(round(number)):,}"


def format_decimal(value: Any, digits: int = 4) -> str:
    number = to_float(value)
    if number is None:
        return "NA"
    return f"{number:.{digits}f}"


def format_pct(value: Any, digits: int = 2) -> str:
    number = to_float(value)
    if number is None:
        return "NA"
    return f"{100.0 * number:.{digits}f}%"


def format_small_pct(value: Any) -> str:
    number = to_float(value)
    if number is None:
        return "NA"
    return f"{100.0 * number:.3f}%"


def format_threshold(value: Any) -> str:
    number = to_float(value)
    if number is None:
        return "NA"
    return f"{number:.6f}"


def find_row(rows: Sequence[dict[str, str]], **criteria: str) -> dict[str, str]:
    for row in rows:
        if all(str(row.get(key)) == str(value) for key, value in criteria.items()):
            return row
    raise KeyError(f"Missing row for criteria={criteria}")


def sort_by_float(rows: Iterable[dict[str, str]], key: str) -> list[dict[str, str]]:
    return sorted(rows, key=lambda row: to_float(row.get(key)) or float("inf"))


def stable_group_ranges(rows: Sequence[dict[str, str]]) -> dict[str, dict[str, Any]]:
    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        grouped[row["candidate_group_key"]].append(row)
    out: dict[str, dict[str, Any]] = {}
    for group, group_rows in grouped.items():
        reasons = [row["primary_fail_reason"] for row in group_rows]
        counts: dict[str, int] = defaultdict(int)
        for reason in reasons:
            counts[reason] += 1
        primary_reason = max(counts.items(), key=lambda item: item[1])[0]
        gate_pass_all = all(str(row["global_gate_pass"]).lower() == "true" for row in group_rows)
        out[group] = {
            "candidate_group_key": group,
            "num_rows": len(group_rows),
            "all_seed_gate_pass": gate_pass_all,
            "min_clean_fpr": min(to_float(row["worst_clean_fpr"]) or 0.0 for row in group_rows),
            "max_clean_fpr": max(to_float(row["worst_clean_fpr"]) or 0.0 for row in group_rows),
            "min_attacked_benign_fpr": min(to_float(row["worst_attacked_benign_fpr"]) or 0.0 for row in group_rows),
            "max_attacked_benign_fpr": max(to_float(row["worst_attacked_benign_fpr"]) or 0.0 for row in group_rows),
            "min_adv_malicious_recall": min(to_float(row["worst_adv_malicious_recall"]) or 0.0 for row in group_rows),
            "max_adv_malicious_recall": max(to_float(row["worst_adv_malicious_recall"]) or 0.0 for row in group_rows),
            "primary_fail_reason": primary_reason,
        }
    return out


def markdown_escape(text: Any) -> str:
    return str(text).replace("|", "\\|")


def markdown_table(columns: Sequence[str], rows: Sequence[Sequence[str]]) -> str:
    lines = [
        "| " + " | ".join(markdown_escape(col) for col in columns) + " |",
        "| " + " | ".join(["---"] * len(columns)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(markdown_escape(cell) for cell in row) + " |")
    return "\n".join(lines)


@dataclass
class TitleBlock:
    title: str
    subtitle: str


@dataclass
class HeadingBlock:
    level: int
    text: str


@dataclass
class ParagraphBlock:
    text: str


@dataclass
class TableBlock:
    title: str
    columns: list[str]
    rows: list[list[str]]
    note: str | None = None


Block = TitleBlock | HeadingBlock | ParagraphBlock | TableBlock


eda_class_balance = read_csv(REPORTS / "eda_advanced_20260305_231027" / "tables" / "class_balance.csv")
eda_protocol_counts = read_csv(REPORTS / "eda_advanced_20260305_231027" / "tables" / "protocol_label_counts.csv")
eda_protocol_shift = read_csv(REPORTS / "eda_advanced_20260305_231027" / "tables" / "train_test_protocol_shift.csv")
eda_attack_family_counts = read_csv(REPORTS / "eda_advanced_20260305_231027" / "tables" / "attack_family_counts.csv")
baseline_metrics = read_csv(REPORTS / "baseline_models_stdlib_20260305_234858" / "metrics_summary.csv")
full_gpu_metrics = read_csv(REPORTS / "full_gpu_models_20260306_001638" / "metrics_summary.csv")
full_gpu_protocol = read_csv(REPORTS / "full_gpu_models_20260306_001638" / "slice_metrics_protocol.csv")
hpo_metrics = read_csv(REPORTS / "full_gpu_hpo_models_20260306_134806" / "metrics_summary.csv")
routed_metrics = read_csv(REPORTS / "full_gpu_hpo_models_20260306_153556" / "metrics_summary.csv")
routed_protocol_models = read_csv(REPORTS / "full_gpu_hpo_models_20260306_153556" / "metrics_summary_per_protocol_models.csv")
leakguard_metrics = read_csv(REPORTS / "full_gpu_hpo_models_20260306_195851" / "metrics_summary.csv")
leakguard_protocol_models = read_csv(REPORTS / "full_gpu_hpo_models_20260306_195851" / "metrics_summary_per_protocol_models.csv")
explainability_global = read_csv(REPORTS / "full_gpu_hpo_models_20260306_195851" / "xgb_explainability" / "global_feature_importance.csv")
robust_protocol = read_csv(REPORTS / "full_gpu_hpo_models_20260306_195851" / "xgb_robustness_realistic_full_20260308_212054" / "robustness_query_metrics_protocol.csv")
wifi_hardening_103936 = read_json(REPORTS / "full_gpu_hpo_models_20260306_195851_wifi_robust_v1_20260309_103936" / "hardening_summary.json")
wifi_hardening_135449 = read_json(REPORTS / "full_gpu_hpo_models_20260306_195851_wifi_robust_v1_20260309_135449" / "hardening_summary.json")
wifi_hardening_180250 = read_json(REPORTS / "full_gpu_hpo_models_20260306_195851_wifi_robust_v1_20260309_180250" / "hardening_summary.json")
wifi_rebalance = read_csv(REPORTS / "full_gpu_hpo_models_20260306_195851_wifi_rebalance_matrix_v1_20260309_204053" / "decision_table.csv")
data_audit_summary = read_json(REPORTS / "data_audit_20260313_124851" / "summary.json")
data_cap_summary = read_csv(
    REPORTS / "full_gpu_hpo_models_20260306_195851_protocol_multimodel_robust_matrix_v1_20260314_112105" / "data_cap_summary.csv"
)
decision_table_global = read_csv(
    REPORTS / "full_gpu_hpo_models_20260306_195851_protocol_multimodel_robust_matrix_v1_20260314_112105" / "decision_table_global.csv"
)
stability_consistency = read_csv(
    REPORTS / "full_gpu_hpo_models_20260306_195851_protocol_multimodel_robust_matrix_v1_20260314_112105" / "stability_consistency_summary.csv"
)
stability_check = read_csv(
    REPORTS / "full_gpu_hpo_models_20260306_195851_protocol_multimodel_robust_matrix_v1_20260314_112105" / "stability_check_global.csv"
)
final_test_metrics = read_csv(
    REPORTS
    / "full_gpu_hpo_models_20260306_195851_protocol_multimodel_robust_matrix_v1_20260314_112105"
    / "catboost_E_test_predictions_metrics_with_thresholds.csv"
)
matrix_summary = read_json(
    REPORTS / "full_gpu_hpo_models_20260306_195851_protocol_multimodel_robust_matrix_v1_20260314_112105" / "matrix_summary.json"
)
escalation = read_json(
    REPORTS / "full_gpu_hpo_models_20260306_195851_protocol_multimodel_robust_matrix_v1_20260314_112105" / "escalation_recommendation.json"
)
stable_near_final = read_csv(
    REPORTS / "full_gpu_hpo_models_20260306_195851_protocol_multimodel_robust_matrix_v1_20260314_003108" / "stability_consistency_summary.csv"
)
external_benign = read_csv(
    REPORTS
    / "full_gpu_hpo_models_20260306_195851_protocol_multimodel_robust_matrix_v1_20260314_112105"
    / "external_benign_data_summary.csv"
)

baseline_logistic = find_row(baseline_metrics, model="logistic_sgd")
baseline_mlp = find_row(baseline_metrics, model="mlp_1hidden")
baseline_adaboost = find_row(baseline_metrics, model="adaboost_stumps")
baseline_stump = find_row(baseline_metrics, model="tree_stump")

full_gpu_mlp = find_row(full_gpu_metrics, model="complex_mlp")
full_gpu_catboost = find_row(full_gpu_metrics, model="catboost")
full_gpu_xgboost = find_row(full_gpu_metrics, model="xgboost")

hpo_mlp = find_row(hpo_metrics, model="complex_mlp_tuned")
hpo_xgboost = find_row(hpo_metrics, model="xgboost_tuned")

routed_xgboost = find_row(routed_metrics, model="xgboost_tuned__protocol_routed")
leakguard_xgboost = find_row(leakguard_metrics, model="xgboost_tuned__protocol_routed")
leakguard_ensemble = find_row(leakguard_metrics, model="weighted_ensemble__protocol_routed")

full_gpu_bt = find_row(full_gpu_protocol, model="complex_mlp", protocol_hint="bluetooth")
full_gpu_wifi = find_row(full_gpu_protocol, model="complex_mlp", protocol_hint="wifi")
routed_bt = find_row(routed_protocol_models, model="xgboost_tuned__bluetooth", protocol_hint="bluetooth")
routed_wifi = find_row(routed_protocol_models, model="xgboost_tuned__wifi", protocol_hint="wifi")
leakguard_bt = find_row(leakguard_protocol_models, model="xgboost_tuned__bluetooth", protocol_hint="bluetooth")
leakguard_wifi = find_row(leakguard_protocol_models, model="xgboost_tuned__wifi", protocol_hint="wifi")

protocol_totals: dict[str, dict[str, Any]] = defaultdict(dict)
for row in eda_protocol_counts:
    protocol = row["protocol_hint"]
    protocol_totals[protocol]["rows"] = to_float(row["rows"]) or 0.0
    protocol_totals[protocol]["benign_rows"] = to_float(row["benign_rows"]) or 0.0
    protocol_totals[protocol]["attack_rows"] = to_float(row["attack_rows"]) or 0.0
    protocol_totals[protocol]["benign_ratio"] = to_float(row["benign_ratio"]) or 0.0
    protocol_totals[protocol]["attack_ratio"] = to_float(row["attack_ratio"]) or 0.0
for row in eda_protocol_shift:
    protocol = row["protocol_hint"]
    protocol_totals[protocol]["train_rows"] = to_float(row["train_rows"]) or 0.0
    protocol_totals[protocol]["test_rows"] = to_float(row["test_rows"]) or 0.0
    protocol_totals[protocol]["train_share"] = to_float(row["train_share"]) or 0.0
    protocol_totals[protocol]["test_share"] = to_float(row["test_share"]) or 0.0

attack_count = find_row(eda_class_balance, label="attack")
benign_count = find_row(eda_class_balance, label="benign")
top_attack_family = [row for row in sort_by_float(eda_attack_family_counts, "rows")[::-1] if row["attack_family"] != "Benign"][:3]

top_features_by_protocol: dict[str, list[str]] = defaultdict(list)
for row in sort_by_float(explainability_global, "rank"):
    protocol = row["protocol_hint"]
    if len(top_features_by_protocol[protocol]) < 3:
        top_features_by_protocol[protocol].append(row["feature"])

wifi_benign_robust_rows = [
    row
    for row in robust_protocol
    if row["attack_method"] == "query_sparse_hillclimb_benign" and row["protocol_hint"] == "wifi"
]
wifi_benign_robust_rows = sorted(
    wifi_benign_robust_rows,
    key=lambda row: to_float(row["epsilon"]) or 0.0,
)

hardening_runs = [
    ("20260309_103936", "Conservative FPR-first hardening", wifi_hardening_103936, "Pre-adversarial-recall gate"),
    ("20260309_135449", "Constraint-aware pilot", wifi_hardening_135449, "Adds adversarial malicious recall control"),
    ("20260309_180250", "Full-budget hardening", wifi_hardening_180250, "Best malicious-side recall, but too many false alarms"),
]

wifi_rebalance_a = find_row(wifi_rebalance, family_id="A")

stability_ranges = stable_group_ranges(stability_check)
stability_lookup = {row["candidate_group_key"]: row for row in stability_consistency}
decision_lookup = {row["candidate_group_key"]: row for row in decision_table_global}

near_final_stable_groups = [
    row["candidate_group_key"]
    for row in stable_near_final
    if str(row["consistent_gate_pass"]).lower() == "true"
]
final_stable_groups = [
    row["candidate_group_key"]
    for row in stability_consistency
    if str(row["consistent_gate_pass"]).lower() == "true"
]

catboost_e_range = stability_ranges["catboost__E"]

protocol_test_lookup = {row["scope"]: row for row in final_test_metrics}
wifi_test = protocol_test_lookup["wifi"]
mqtt_test = protocol_test_lookup["mqtt"]
bluetooth_test = protocol_test_lookup["bluetooth"]
global_test = protocol_test_lookup["global_protocol_routed"]

bluetooth_benign_cap = find_row(data_cap_summary, protocol="bluetooth", label="0")
mqtt_benign_cap = find_row(data_cap_summary, protocol="mqtt", label="0")
wifi_benign_cap = find_row(data_cap_summary, protocol="wifi", label="0")


def build_blocks() -> list[Block]:
    blocks: list[Block] = []
    blocks.append(
        TitleBlock(
            title="Experimental Development and Final Selection of the Flow-Based IoMT IDS",
            subtitle=(
                "Script-verified thesis chapter reconstructed from saved results, chronological context, "
                "and reverse-engineered training and evaluation code"
            ),
        )
    )

    chronology_rows = [
        [
            "2026-03-05",
            "Dataset merge and advanced EDA",
            "merge_ciciomt_with_metadata.py; generate_advanced_eda_report.py",
            "reports/eda_advanced_20260305_231027",
            "Establish whether the problem is learnable and whether protocol slices must be treated separately.",
        ],
        [
            "2026-03-05",
            "Deterministic reduced-sample baselines",
            "train_baseline_models_stdlib.py",
            "reports/baseline_models_stdlib_20260305_234858",
            "Spend GPU budget only after proving the feature space already carries usable signal.",
        ],
        [
            "2026-03-06",
            "Full-data GPU clean modeling",
            "train_full_gpu_models.py",
            "reports/full_gpu_models_20260306_001638",
            "Measure how much performance survives when the whole dataset and stricter thresholding are used.",
        ],
        [
            "2026-03-06",
            "Constrained GPU HPO",
            "train_hpo_gpu_models.py",
            "reports/full_gpu_hpo_models_20260306_134806",
            "Optimize under low-FPR intent rather than raw clean F1.",
        ],
        [
            "2026-03-06",
            "Protocol-routed HPO",
            "train_hpo_gpu_models.py",
            "reports/full_gpu_hpo_models_20260306_153556",
            "Split the problem by protocol after the global model still looked operationally uneven.",
        ],
        [
            "2026-03-06",
            "Split repair and leakguard",
            "train_hpo_gpu_models_fprfix.py; train_hpo_gpu_models_leakguard.py",
            "reports/full_gpu_hpo_models_20260306_172441 -> 195851",
            "Convert impressive routed metrics into trustworthy metrics by hardening validation and leakage controls.",
        ],
        [
            "2026-03-07 to 2026-03-09",
            "Explainability and realistic robustness",
            "xgb_protocol_ids_utils.py; generate_xgb_explainability_artifacts.py; evaluate_xgb_robustness.py",
            "reports/full_gpu_hpo_models_20260306_195851/xgb_explainability and xgb_robustness_realistic_full_20260308_212054",
            "Determine whether the clean low-FPR model remains credible under constrained feature-space attack.",
        ],
        [
            "2026-03-09",
            "WiFi hardening and rebalance",
            "train_wifi_robust_hardening.py; train_wifi_robust_rebalance_matrix.py",
            "reports/...wifi_robust_v1_* and ...wifi_rebalance_matrix_v1_20260309_204053",
            "Repair the benign-side WiFi weakness revealed by robustness evaluation.",
        ],
        [
            "2026-03-10 to 2026-03-14",
            "Protocol-wide robust matrix and stability reruns",
            "train_protocol_multimodel_robust_matrix.py",
            "reports/...protocol_multimodel_robust_matrix_v1_20260314_112105",
            "Select the deployment baseline by gate passing, stability, and saved-artifact availability rather than by coarse rank alone.",
        ],
    ]

    blocks.append(HeadingBlock(1, "1. Chapter Aim and Evidence Base"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                This chapter turns the experiment log into a defensible thesis narrative. The goal is not to repeat every
                artifact that exists under reports/. The goal is to explain, in chronological order, why each experimental
                decision was taken, what the saved results revealed at that point, and why those results forced the next
                step. The evidence base for the chapter is therefore deliberately narrower than the total artifact set:
                only the tables that changed the direction of the project are retained in the main body, while the finer
                implementation details are pushed into prose and the script-verification appendix.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                Three sources are treated as authoritative. The first is the chronological context log, which records the
                reasoning and the sequence of pivots. The second is the result folders under reports/, which provide the
                saved metrics and thresholds used at the time. The third is the Python code itself. I re-read the merge,
                training, thresholding, robustness, and matrix-selection scripts so that this document reflects what the
                pipeline actually did rather than what would be convenient to claim afterwards.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 1. Chronological structure of the experimental campaign",
            columns=["Date", "Phase", "Verified script(s)", "Authoritative artifact", "Decision carried forward"],
            rows=chronology_rows,
            note=(
                "Only milestones that materially changed the research direction are kept in the chapter body. "
                "Appendix A maps the scripts to the exact methodological consequences in more detail."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The chronology in Table 1 also clarifies the scope of the contribution. This project did not move in a
                straight line from one clean leaderboard to the next. It began as a learnability question, became a
                protocol-heterogeneity question, then became a robustness question, and finally became a stability and
                artifact-selection question. That sequence matters because the final model would have been different if
                the campaign had stopped after any earlier stage.
                """
            )
        )
    )

    protocol_rows: list[list[str]] = []
    for protocol in ["wifi", "mqtt", "bluetooth"]:
        row = protocol_totals[protocol]
        protocol_rows.append(
            [
                protocol,
                format_count(row["rows"]),
                format_count(row["benign_rows"]),
                format_count(row["attack_rows"]),
                format_pct(row["attack_ratio"]),
                format_pct(row["train_share"]),
                format_pct(row["test_share"]),
            ]
        )

    blocks.append(HeadingBlock(1, "2. Dataset Construction and EDA Defined the Real Problem"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The first non-trivial step was not model training but dataset construction. The script
                merge_ciciomt_with_metadata.py merged the CICIoMT CSV files into metadata_train.csv and metadata_test.csv
                while deriving stable metadata fields such as source_relpath, source_modality, protocol_hint, device,
                scenario, attack_name, attack_family, label, and source_row_index. Attack captures preserved folder-based
                train/test assignment when it existed. Profiling traffic without explicit split folders was assigned by
                deterministic hashing. This design mattered because every later claim about protocol-aware training,
                family-wise slicing, leakage prevention, and routed inference depends on those metadata columns being
                correct and reproducible.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                Once the merged tables existed, the advanced EDA immediately showed that the classification task was not a
                clean, balanced binary benchmark. The dataset contained {format_count(attack_count['rows'])} attack rows
                versus {format_count(benign_count['rows'])} benign rows, meaning the positive class represented
                {format_pct(attack_count['share'])} of all observations. The dominant attack mass was concentrated in the
                {top_attack_family[0]['attack_family']} family, followed by {top_attack_family[1]['attack_family']} and
                {top_attack_family[2]['attack_family']}. In other words, a high global F1 score would be easy to inflate
                unless the analysis paid attention to false positives, minority protocols, and family concentration.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 2. Protocol composition and train-test share after metadata merge",
            columns=[
                "Protocol",
                "Rows",
                "Benign rows",
                "Attack rows",
                "Attack share",
                "Train share",
                "Test share",
            ],
            rows=protocol_rows,
            note=(
                "The large WiFi share explains why global metrics alone were never sufficient. A protocol-specific failure "
                "could remain operationally serious while barely moving the overall score."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                Table 2 made the central design pressure visible very early. WiFi dominated the corpus, MQTT was much
                smaller, and Bluetooth was both the smallest and structurally different slice. That asymmetry implied two
                things. First, low-FPR thresholding would matter at least as much as raw recall because even a small false
                positive rate on the dominant benign slice would be costly in practice. Second, protocol-specific analysis
                was not an optional diagnostic; it was the correct problem formulation. The EDA also confirmed strong
                feature signal in flow statistics such as flag counts, variance-derived measures, and volume features, so
                the flow-only approach was worth pursuing without payload inspection.
                """
            )
        )
    )

    baseline_rows = [
        [
            "logistic_sgd",
            format_threshold(baseline_logistic["threshold"]),
            format_decimal(baseline_logistic["precision"]),
            format_decimal(baseline_logistic["recall"]),
            format_decimal(baseline_logistic["f1"]),
            format_small_pct(baseline_logistic["fpr"]),
            format_decimal(baseline_logistic["roc_auc"]),
        ],
        [
            "mlp_1hidden",
            format_threshold(baseline_mlp["threshold"]),
            format_decimal(baseline_mlp["precision"]),
            format_decimal(baseline_mlp["recall"]),
            format_decimal(baseline_mlp["f1"]),
            format_small_pct(baseline_mlp["fpr"]),
            format_decimal(baseline_mlp["roc_auc"]),
        ],
        [
            "adaboost_stumps",
            format_threshold(baseline_adaboost["threshold"]),
            format_decimal(baseline_adaboost["precision"]),
            format_decimal(baseline_adaboost["recall"]),
            format_decimal(baseline_adaboost["f1"]),
            format_small_pct(baseline_adaboost["fpr"]),
            format_decimal(baseline_adaboost["roc_auc"]),
        ],
        [
            "tree_stump",
            format_threshold(baseline_stump["threshold"]),
            format_decimal(baseline_stump["precision"]),
            format_decimal(baseline_stump["recall"]),
            format_decimal(baseline_stump["f1"]),
            format_small_pct(baseline_stump["fpr"]),
            format_decimal(baseline_stump["roc_auc"]),
        ],
    ]

    blocks.append(HeadingBlock(1, "3. Reduced-Sample Baselines Proved Learnability Before GPU Scaling"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The project did not jump directly from EDA into expensive GPU experiments. The script
                train_baseline_models_stdlib.py was written specifically to establish a cheap and deterministic
                learnability check. It used stable hashing of source_relpath and source_row_index to route a reproducible
                subset of rows into validation, reservoir sampling to cap train/validation/test volumes by class, and
                training-time standardization computed only from the sampled training set. That design kept the baseline
                stage lightweight without relaxing the separation discipline needed later.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 3. Reduced-sample baseline leaderboard",
            columns=["Model", "Threshold", "Precision", "Recall", "F1", "FPR", "ROC-AUC"],
            rows=baseline_rows,
            note="These results came from the reduced deterministic sample, not from the full dataset.",
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                Even on the reduced sample, the signal was already clear. The logistic baseline reached an F1 of
                {format_decimal(baseline_logistic['f1'])} with an FPR of {format_small_pct(baseline_logistic['fpr'])},
                while the single stump model collapsed to an F1 of {format_decimal(baseline_stump['f1'])} and an FPR of
                {format_pct(baseline_stump['fpr'])}. That contrast mattered more than the exact ordering. It showed that
                the feature space was informative enough that stronger families and full-data training were worth the
                compute, but it also warned that overly shallow models would not survive protocol heterogeneity.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                This stage therefore answered the first research question: yes, the flow features carried real IDS signal.
                It did not answer the deployment question. The baseline sample was too small, too clean, and too detached
                from the eventual operating point. The next step was therefore not to tune the baseline harder, but to
                move to full-data GPU models under an explicit threshold policy.
                """
            )
        )
    )

    clean_progression_rows = [
        [
            "Reduced deterministic sample",
            "logistic_sgd",
            format_decimal(baseline_logistic["f1"]),
            format_small_pct(baseline_logistic["fpr"]),
            format_decimal(baseline_logistic["roc_auc"]),
            "Cheap proof that learnability exists",
        ],
        [
            "Full-data GPU clean run",
            "complex_mlp",
            format_decimal(full_gpu_mlp["f1"]),
            format_pct(full_gpu_mlp["fpr"]),
            format_decimal(full_gpu_mlp["roc_auc"]),
            "Best clean global score, but still too many false positives",
        ],
        [
            "First constrained GPU HPO",
            "complex_mlp_tuned",
            format_decimal(hpo_mlp["f1"]),
            format_pct(hpo_mlp["fpr"]),
            format_decimal(hpo_mlp["roc_auc"]),
            "Optimization improved, deployment credibility did not",
        ],
        [
            "First constrained GPU HPO",
            "xgboost_tuned",
            format_decimal(hpo_xgboost["f1"]),
            format_pct(hpo_xgboost["fpr"]),
            format_decimal(hpo_xgboost["roc_auc"]),
            "Low-FPR objective on validation still did not prevent test-side FPR drift",
        ],
    ]

    blocks.append(HeadingBlock(1, "4. Full-Data GPU Modeling Improved Scores but Exposed an Operational Weakness"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The next stage was implemented in train_full_gpu_models.py. That script used a deterministic validation
                split and selected the operating threshold by maximizing validation recall subject to FPR <= 1%. This is
                an important detail because it means the early GPU runs were already trying to behave like an IDS, not like
                a generic classifier. The test metrics therefore revealed something scientifically useful: even when the
                validation threshold was chosen under a one-percent false-positive ceiling, the deployed test-side FPR was
                still materially higher.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 4. Clean-model progression before protocol routing",
            columns=["Stage", "Representative model", "F1", "FPR", "ROC-AUC", "Interpretation"],
            rows=clean_progression_rows,
            note=(
                "The first two GPU stages improved headline metrics, but neither eliminated the gap between low-FPR intent "
                "on validation and operational false positives on test."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                Table 4 explains why the thesis could not end at the first strong GPU result. The full-data complex MLP
                reached an F1 of {format_decimal(full_gpu_mlp['f1'])} and an ROC-AUC of
                {format_decimal(full_gpu_mlp['roc_auc'])}, but its test FPR was still {format_pct(full_gpu_mlp['fpr'])}.
                The first GPU HPO stage, implemented in train_hpo_gpu_models.py, made the objective more sophisticated:
                it optimized validation F1 plus a PR-AUC bonus and penalized gaps above the target FPR. Yet the best tuned
                MLP still landed at {format_pct(hpo_mlp['fpr'])} test FPR. These are strong clean classification results,
                but they are not yet a deployment-grade operating point.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                More importantly, the protocol slices showed why the global score was misleading. In the first full-data
                clean run, the representative MLP produced a Bluetooth FPR of {format_pct(full_gpu_bt['fpr'])} while the
                WiFi slice sat at {format_pct(full_gpu_wifi['fpr'])}. A single thresholded global score could therefore
                look excellent while still being operationally uneven. That is the point at which protocol routing stopped
                being a nice-to-have extension and became the next necessary experiment.
                """
            )
        )
    )

    routed_comparison_rows = [
        [
            "Full-data clean MLP",
            "complex_mlp",
            format_decimal(full_gpu_mlp["f1"]),
            format_pct(full_gpu_mlp["fpr"]),
            format_pct(full_gpu_bt["fpr"]),
            format_pct(full_gpu_wifi["fpr"]),
            "Global score high, Bluetooth still expensive",
        ],
        [
            "Initial protocol-routed HPO",
            "xgboost_tuned__protocol_routed",
            format_decimal(routed_xgboost["f1"]),
            format_pct(routed_xgboost["fpr"]),
            format_small_pct(routed_bt["fpr"]),
            format_pct(routed_wifi["fpr"]),
            "Routing removes most protocol-specific pain",
        ],
        [
            "Leakguard routed HPO",
            "xgboost_tuned__protocol_routed",
            format_decimal(leakguard_xgboost["f1"]),
            format_small_pct(leakguard_xgboost["fpr"]),
            format_small_pct(leakguard_bt["fpr"]),
            format_small_pct(leakguard_wifi["fpr"]),
            "Routing survives stronger split controls and leakage checks",
        ],
    ]

    blocks.append(HeadingBlock(1, "5. Protocol Routing Was the First Major Structural Breakthrough"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                Protocol routing was introduced because the earlier results suggested that WiFi, MQTT, and Bluetooth were
                not simply three interchangeable subsets of one traffic distribution. The HPO code already supported
                per-protocol training and thresholding, so the next logical experiment was to let each protocol keep its
                own model and threshold. When this was first executed in the 20260306_153556 run, the improvement was not
                marginal; it changed the shape of the problem.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 5. Why protocol routing mattered",
            columns=["Stage", "Model", "Global F1", "Global FPR", "Bluetooth FPR", "WiFi FPR", "What changed"],
            rows=routed_comparison_rows,
            note=(
                "Bluetooth and WiFi are shown because they carried the decisive operational contrast. MQTT was already "
                "strong in the clean stages and did not drive the routing decision."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                The magnitude of the change is what forced the next methodological reaction. Bluetooth false positives
                dropped from {format_pct(full_gpu_bt['fpr'])} in the clean MLP stage to {format_small_pct(routed_bt['fpr'])}
                under routed XGBoost. WiFi also improved materially. Those gains were too large to accept uncritically. In
                IDS work, a difficult slice becoming almost perfect is exactly the point where one has to check whether the
                validation protocol is flattering the model. The correct scientific response was therefore skepticism, not
                celebration.
                """
            )
        )
    )

    blocks.append(HeadingBlock(1, "6. Split Repair and Leakguard Turned Good Numbers into Trustworthy Numbers"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The suspicion raised by the first routed results directly motivated train_hpo_gpu_models_fprfix.py. That
                script changed the validation construction from a simpler hashed split into a deterministic
                protocol-and-class-aware split with file-level grouping as the primary behavior. Only when a stratum
                collapsed to a single source file was row-level fallback allowed. This mattered because it reduced the
                chance that a tiny or structurally narrow protocol slice would look cleaner than it really was.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The next script, train_hpo_gpu_models_leakguard.py, tightened the pipeline further. It refused to rely on
                positional feature slicing unless the first twenty metadata columns matched the expected prefix exactly. It
                explicitly banned metadata and label-like fields from entering the feature set. It also checked that the
                train and test source_relpath sets were disjoint so that the same capture file could not silently appear on
                both sides. This stage is the methodological trust anchor of the whole campaign, because after it, strong
                results could be interpreted as model behavior rather than as schema luck.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                The leakguard run left two important signals. First, the routed XGBoost system still achieved an F1 of
                {format_decimal(leakguard_xgboost['f1'])} while pushing global FPR down to
                {format_small_pct(leakguard_xgboost['fpr'])}. Second, the weighted ensemble obtained the highest
                validation objective but a much higher test FPR of {format_pct(leakguard_ensemble['fpr'])}. That split in
                behavior is why the project used the routed XGBoost stack as the base platform for explainability and
                robustness. It was not simply the clean winner; it was the clean low-FPR model whose behavior was easiest
                to audit and extend.
                """
            )
        )
    )

    blocks.append(HeadingBlock(1, "7. Explainability Was Used as an Audit Tool, Not Cosmetic Decoration"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                Once the leakguard-routed XGBoost pipeline existed, the next question was whether its decisions could be
                inspected in a way that supported a healthcare IDS argument. The explainability utilities in
                xgb_protocol_ids_utils.py and generate_xgb_explainability_artifacts.py did two useful things. They exposed
                routed predictions with protocol-specific thresholds, and they saved both global and local contribution
                artifacts for the representative true-positive, true-negative, false-positive, and false-negative cases.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                The resulting feature rankings were reassuring. The highest-contribution features were not leaked metadata
                fields, because leakguard had already forbidden those. Instead, they were plausible flow aggregates such as
                {top_features_by_protocol['bluetooth'][0]}, {top_features_by_protocol['bluetooth'][1]}, and
                {top_features_by_protocol['bluetooth'][2]} on Bluetooth, together with analogous traffic-volume and rate
                features on WiFi and MQTT. This did not prove causal correctness, but it did show that the model was
                building its decisions from network-flow behavior rather than from accidental identifiers.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                More importantly, the explainability stage prepared the project for the next pivot. Because the routed
                model could now return per-protocol predictions and local contributions, it became feasible to attack the
                system in a controlled way and then interpret where the model was brittle. The interpretability work was
                therefore not a presentation layer. It was part of the transition from clean evaluation to adversarial
                evaluation.
                """
            )
        )
    )

    robust_rows = []
    for row in wifi_benign_robust_rows:
        robust_rows.append(
            [
                format_decimal(row["epsilon"], 2),
                format_pct(row["fpr"]),
                format_decimal(row["precision"]),
                format_decimal(row["f1"]),
                format_decimal(row["delta_f1"]),
                format_decimal(row["queries_mean"], 1),
            ]
        )

    blocks.append(HeadingBlock(1, "8. Realistic Robustness Changed the Definition of Success"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The robustness stage was where the thesis changed character. The script evaluate_xgb_robustness.py did not
                apply unconstrained adversarial noise. It locked semantic features, respected lower and upper bounds, and
                enforced relation-aware realism constraints before scoring candidate perturbations. It also attacked both
                malicious and benign traffic. That second choice turned out to be decisive. Malicious-side evasion under
                the realistic query attack degraded performance only modestly. Benign-side perturbation on WiFi, however,
                revealed a much more serious operational weakness.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 6. WiFi benign-side degradation under realistic query attack",
            columns=["Epsilon", "Attacked benign FPR", "Precision", "F1", "Delta F1", "Mean queries"],
            rows=robust_rows,
            note=(
                "All rows come from query_sparse_hillclimb_benign on the WiFi protocol slice of the leakguard-routed "
                "XGBoost system."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                Table 6 is the single most important turning-point table in the project. At epsilon 0.01, attacked benign
                FPR jumped to {format_pct(wifi_benign_robust_rows[1]['fpr'])}. By epsilon 0.10, it reached
                {format_pct(wifi_benign_robust_rows[-1]['fpr'])}, and the corresponding F1 dropped by
                {abs((to_float(wifi_benign_robust_rows[-1]['delta_f1']) or 0.0)):.4f}. This result forced a conceptual
                shift. Clean low-FPR performance was no longer enough. The real deployment question became whether the
                model could avoid being manipulated into false alarms on benign traffic while still keeping malicious
                recall high.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                That observation explains why the rest of the project is dominated by hardening, rebalance, gates, and
                stability rather than by conventional clean leaderboard tuning. Once benign-side attack sensitivity became
                visible, the correct objective was not to maximize one more decimal place of clean F1. It was to define an
                acceptable operating envelope and select only the candidates that could remain inside it.
                """
            )
        )
    )

    hardening_rows: list[list[str]] = []
    for run_id, label, summary, verdict in hardening_runs:
        selected = summary["thresholds"]["selection"]["selected_metrics"]
        hardening_rows.append(
            [
                run_id,
                label,
                format_threshold(summary["thresholds"]["wifi_new"]),
                format_pct(selected["clean_fpr"]),
                format_pct(selected["attacked_benign_fpr"]),
                format_decimal(selected.get("attacked_malicious_recall")),
                format_decimal(selected["clean_f1"]),
                verdict,
            ]
        )
    hardening_rows.append(
        [
            "20260309_204053",
            "WiFi rebalance winner (family A)",
            format_threshold(wifi_rebalance_a["selected_threshold"]),
            format_pct(wifi_rebalance_a["clean_fpr"]),
            format_pct(wifi_rebalance_a["attacked_benign_fpr"]),
            format_decimal(wifi_rebalance_a["adv_malicious_recall"]),
            format_decimal(wifi_rebalance_a["clean_f1"]),
            "Only WiFi candidate that passed all explicit gates",
        ]
    )

    blocks.append(HeadingBlock(1, "9. WiFi Hardening and Rebalance Converted a Weakness into a Formal Selection Problem"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                Because the benign-side failure was concentrated on WiFi, the next experiments became deliberately
                WiFi-specific. The script train_wifi_robust_hardening.py generated adversarial benign hard negatives,
                selected a new threshold under clean-recall and optional adversarial-recall constraints, and measured the
                trade-off between benign protection and malicious detection. These runs were not simply alternative
                retrains. They were controlled attempts to discover which kind of hardening pressure produced a useful
                operating point.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 7. WiFi hardening sequence and the rebalance winner",
            columns=[
                "Run",
                "Configuration",
                "Selected threshold",
                "Clean FPR",
                "Attacked benign FPR",
                "Adv. malicious recall",
                "Clean F1",
                "Interpretation",
            ],
            rows=hardening_rows,
            note=(
                "The first run focused on false positives only. The second and third runs introduced adversarial-malicious "
                "constraints. The rebalance run then reframed the decision as a family-based gate check."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                The WiFi sequence is a useful example of why more training budget does not automatically mean a better
                deployment point. The conservative run pushed attacked benign FPR down to {format_pct(wifi_hardening_103936['thresholds']['selection']['selected_metrics']['attacked_benign_fpr'])}
                while keeping clean FPR at {format_pct(wifi_hardening_103936['thresholds']['selection']['selected_metrics']['clean_fpr'])},
                but it was not yet checking malicious-side robustness. The constraint-aware pilot was more balanced and
                even achieved adversarial malicious recall of {format_decimal(wifi_hardening_135449['thresholds']['selection']['selected_metrics']['attacked_malicious_recall'])},
                yet its clean FPR climbed to {format_pct(wifi_hardening_135449['thresholds']['selection']['selected_metrics']['clean_fpr'])}.
                The full-budget hardening run looked attractive from the malicious side, but it overshot badly on benign
                false alarms, landing at roughly {format_pct(wifi_hardening_180250['thresholds']['selection']['selected_metrics']['clean_fpr'])}
                clean FPR and {format_pct(wifi_hardening_180250['thresholds']['selection']['selected_metrics']['attacked_benign_fpr'])}
                attacked benign FPR.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                This is why train_wifi_robust_rebalance_matrix.py mattered. It stopped treating the problem as one
                threshold on one retrain and instead evaluated model families against explicit gates. Family A passed with
                clean FPR {format_pct(wifi_rebalance_a['clean_fpr'])}, attacked benign FPR
                {format_pct(wifi_rebalance_a['attacked_benign_fpr'])}, and adversarial malicious recall
                {format_decimal(wifi_rebalance_a['adv_malicious_recall'])}. Families B and C preserved very low FPR but
                collapsed on adversarial malicious recall. The next decision therefore became obvious: the WiFi gate-based
                logic had to be promoted to a protocol-wide robust matrix.
                """
            )
        )
    )

    gate_rows = [
        ["Clean FPR max", format_pct(matrix_summary["gates"]["clean_fpr_max"]), "Upper bound for clean false positives"],
        [
            "Attacked benign FPR max",
            format_pct(matrix_summary["gates"]["attacked_benign_fpr_max"]),
            "Upper bound for benign-side adversarial false positives",
        ],
        [
            "Adversarial malicious recall min",
            format_pct(matrix_summary["gates"]["adv_malicious_recall_min"]),
            "Lower bound for malicious recall after attack",
        ],
    ]

    blocks.append(HeadingBlock(1, "10. The Protocol-Wide Robust Matrix Reframed Selection Around Gates and Stability"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The final stage was implemented in train_protocol_multimodel_robust_matrix.py. This script is important
                not because it trained more candidates than the earlier code, but because it changed the meaning of model
                selection. The allowed final families were restricted to XGBoost and CatBoost. Each candidate began with a
                clean warm-fit, received a base threshold selected from clean FPR, generated benign hard negatives and
                malicious adversarial examples, refit on the weighted augmented data, and then selected its final threshold
                using a lexicographic gate search. In other words, passing the robust operating constraints was ranked
                ahead of squeezing out the last clean-F1 advantage.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 8. Gates enforced by the final protocol-wide robust matrix",
            columns=["Gate", "Value", "Meaning"],
            rows=gate_rows,
            note=(
                "The matrix used an internal attacked-benign safety margin during threshold search, but the three values "
                "shown here were the public acceptance criteria for candidate promotion."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The family logic also reflects what the project learned along the way. Families A and B are essentially
                precision- and balance-oriented references. Families C, D, and E exist because the search needed targeted
                Bluetooth-recovery behavior without giving back the WiFi false-positive improvements. That is why the final
                answer eventually came from family E. Family E was not an arbitrary later variant; it was created because
                the earlier evidence showed that the bottleneck could no longer be solved by a single generic robust
                retrain.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The matrix configuration in the final authoritative run also makes clear that this was a deliberate search,
                not an improvised sweep. The attack source mode was fixed to the XGBoost base run, the family pack was the
                Bluetooth-recovery fallback pack, coarse search used adaptive-fraction sampling with lightweight query
                budgets, and stability search reused the same split seed while increasing sample support and query budget.
                The project moved from ordinary tuning to controlled search design because that is what the robustness
                evidence demanded.
                """
            )
        )
    )

    caveat_rows = [
        [
            "wifi",
            format_count(wifi_benign_cap["rows"]),
            format_count(wifi_benign_cap["val_rows"]),
            format_pct(wifi_benign_cap["largest_file_share"]),
            wifi_benign_cap["fallback_type"],
            "Main operational benchmark; strong validation support but still dominant in the corpus.",
        ],
        [
            "mqtt",
            format_count(mqtt_benign_cap["rows"]),
            format_count(mqtt_benign_cap["val_rows"]),
            format_pct(mqtt_benign_cap["largest_file_share"]),
            mqtt_benign_cap["fallback_type"],
            "Validation has benign support, but held-out test has no benign negatives, so FPR and ROC-AUC cannot be estimated there.",
        ],
        [
            "bluetooth",
            format_count(bluetooth_benign_cap["rows"]),
            format_count(bluetooth_benign_cap["val_rows"]),
            format_pct(bluetooth_benign_cap["largest_file_share"]),
            bluetooth_benign_cap["fallback_type"],
            f"Benign evidence is highly concentrated; floor repair moved {format_count(bluetooth_benign_cap['rows_moved_to_train'])} rows back into training.",
        ],
    ]

    blocks.append(HeadingBlock(1, "11. Data Audit and Feasibility Limits Were Treated as Results, Not Footnotes"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                The data audit performed on 2026-03-13 is one of the reasons the final chapter can be honest without
                being weak. It confirmed {format_count(data_audit_summary['train_rows'])} training rows,
                {format_count(data_audit_summary['test_rows'])} test rows, and a feature space of
                {format_count(data_audit_summary['feature_count'])} flow variables. It also confirmed several structural
                limits that the later robustness results had already hinted at: WiFi dominance, MQTT test benign absence,
                Bluetooth benign concentration, and many near-constant protocol-specific features. None of those findings
                invalidates the final model, but each one changes how the thesis should phrase its claims.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 9. Validation support and caveats by protocol",
            columns=[
                "Protocol",
                "Benign rows",
                "Validation benign rows",
                "Largest file share",
                "Split fallback",
                "Why it matters",
            ],
            rows=caveat_rows,
            note=(
                "These rows come from the final authoritative robust-matrix run because those are the statistics that "
                "directly governed the final low-FPR gate checks."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                The Bluetooth row in Table 9 is especially important. Before the floor repair, the final matrix had only
                {format_count(bluetooth_benign_cap['benign_train_before_floor'])} Bluetooth benign training rows, which is
                not enough support for stable low-FPR calibration. The pipeline therefore moved
                {format_count(bluetooth_benign_cap['rows_moved_to_train'])} rows back into training to reach a floor of
                {format_count(bluetooth_benign_cap['benign_train_after_floor'])}. That should be presented as explicit
                statistical housekeeping, not as a hidden trick. The MQTT caveat is different: the issue is not missing
                validation support, but the lack of benign negatives in the held-out test set. That is why the final MQTT
                FPR and ROC-AUC entries are honestly left undefined.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                This section also explains why some artifacts are absent from the final claim. External benign augmentation
                stayed disabled in the final authoritative run, so the promoted model should be described as the best
                in-project deployment baseline, not as a domain-generalized IDS across unmatched external datasets. That is
                a tighter claim, but it is also the claim the evidence actually supports.
                """
            )
        )
    )

    stability_table_rows: list[list[str]] = []
    for group in ["xgboost__A", "xgboost__C", "catboost__E"]:
        rank_row = decision_lookup[group]
        range_row = stability_ranges[group]
        consistency_row = stability_lookup[group]
        stability_table_rows.append(
            [
                group,
                rank_row["final_rank"],
                "Yes" if str(consistency_row["consistent_gate_pass"]).lower() == "true" else "No",
                consistency_row["num_seeds_checked"],
                format_pct(range_row["max_clean_fpr"]),
                format_pct(range_row["max_attacked_benign_fpr"]),
                format_pct(range_row["min_adv_malicious_recall"]),
                range_row["primary_fail_reason"],
            ]
        )

    blocks.append(HeadingBlock(1, "12. Stability Reruns Determined the Final Winner"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The last selection problem was no longer about finding a candidate that could pass the gates once. It was
                about finding a candidate that kept passing after the search was rerun, the sampling seed moved, and model
                artifacts were actually saved. This is the stage where a superficial leaderboard reading would have chosen
                the wrong winner. In the final authoritative 20260314_112105 run, the seed-42 global ranking placed
                XGBoost families at the top. But the stability tables tell a different story.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 10. Rank on the main run versus stability across reruns",
            columns=[
                "Candidate group",
                "Seed-42 rank",
                "All-seed gate pass",
                "Seeds checked",
                "Worst clean FPR",
                "Worst attacked benign FPR",
                "Worst adv. recall",
                "Primary failure reason",
            ],
            rows=stability_table_rows,
            note=(
                "Seed-42 rank comes from decision_table_global.csv. Stability comes from the nine-seed consistency and "
                "stability-check outputs of the final authoritative run."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                Table 10 is the most important selection table in the chapter. The rank-1 candidate on the main run was
                xgboost__A, and xgboost__C was rank 2. Neither survived the stability filter. Both repeatedly failed for
                the same reason: FPR drift. By contrast, catboost__E was only rank 6 on the seed-42 decision table, yet it
                was the only candidate that remained gate-pass consistent across all
                {format_count(stability_lookup['catboost__E']['num_seeds_checked'])} checked seeds. Its attacked benign
                FPR remained between {format_pct(catboost_e_range['min_attacked_benign_fpr'])} and
                {format_pct(catboost_e_range['max_attacked_benign_fpr'])}, while adversarial malicious recall stayed
                between {format_pct(catboost_e_range['min_adv_malicious_recall'])} and
                {format_pct(catboost_e_range['max_adv_malicious_recall'])}.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                The March 14 near-final run makes this point even stronger. In the earlier 20260314_003108 pass, the
                consistent groups were {', '.join(near_final_stable_groups)}. After the artifact-persistent rerun, the
                final stable set became {', '.join(final_stable_groups)}. That shift is exactly why the project kept going.
                A candidate that looks stable before artifact persistence is only analytically interesting. A candidate
                that remains stable after artifact persistence and rerun is a deployment baseline. The escalation output of
                the final run captured this directly: recommendation={escalation['recommendation']}, with no protocol hard
                cap triggered.
                """
            )
        )
    )

    final_metric_rows = [
        [
            "WiFi",
            format_threshold(wifi_test["threshold"]),
            format_decimal(wifi_test["precision"]),
            format_decimal(wifi_test["recall"]),
            format_decimal(wifi_test["f1"]),
            format_pct(wifi_test["fpr"]),
            format_count(wifi_test["fp"]),
            format_count(wifi_test["fn"]),
        ],
        [
            "MQTT",
            format_threshold(mqtt_test["threshold"]),
            format_decimal(mqtt_test["precision"]),
            format_decimal(mqtt_test["recall"]),
            format_decimal(mqtt_test["f1"]),
            "NA",
            format_count(mqtt_test["fp"]),
            format_count(mqtt_test["fn"]),
        ],
        [
            "Bluetooth",
            format_threshold(bluetooth_test["threshold"]),
            format_decimal(bluetooth_test["precision"]),
            format_decimal(bluetooth_test["recall"]),
            format_decimal(bluetooth_test["f1"]),
            format_pct(bluetooth_test["fpr"]),
            format_count(bluetooth_test["fp"]),
            format_count(bluetooth_test["fn"]),
        ],
        [
            "Protocol-routed global",
            "protocol-specific",
            format_decimal(global_test["precision"]),
            format_decimal(global_test["recall"]),
            format_decimal(global_test["f1"]),
            format_pct(global_test["fpr"]),
            format_count(global_test["fp"]),
            format_count(global_test["fn"]),
        ],
    ]

    blocks.append(HeadingBlock(1, "13. Final Deployment-Ready Baseline"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                After stability and artifact persistence were enforced, the final deployment baseline was the protocol-routed
                catboost__E system saved by the 20260314_112105 run. The final test metrics were computed by applying the
                saved protocol-specific thresholds to metadata_test.csv and evaluating the routed predictions directly.
                This matters because it means the closing table is not a validation artifact; it is a held-out test
                summary for the actual saved model family that survived the full campaign.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 11. Final held-out test metrics for the promoted catboost__E deployment baseline",
            columns=["Scope", "Threshold", "Precision", "Recall", "F1", "FPR", "FP", "FN"],
            rows=final_metric_rows,
            note=(
                "MQTT FPR and ROC-AUC remain undefined on held-out test because there are no MQTT benign negatives in the "
                "test split."
            ),
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                The promoted baseline achieved global precision {format_decimal(global_test['precision'])}, recall
                {format_decimal(global_test['recall'])}, F1 {format_decimal(global_test['f1'])}, and FPR
                {format_pct(global_test['fpr'])}. On WiFi, only {format_count(wifi_test['fp'])} false positives were
                produced over {format_count(wifi_test['n_rows'])} rows. Bluetooth achieved zero false positives in the
                held-out test slice. MQTT preserved very strong recall and precision, but its benign-side error rates
                cannot be estimated honestly on the test split because no benign negatives exist there.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                f"""
                The right thesis claim is therefore precise: catboost__E is the best deployment-ready in-project baseline
                produced by this campaign, not because it was the flashiest coarse winner, but because it is the only saved
                candidate that survived the complete chain of evidence. That chain included metadata-aware data assembly,
                protocol-aware evaluation, leakguard validation, realistic robustness, WiFi hardening, protocol-wide robust
                gates, multi-seed stability, and explicit artifact persistence. The fact that external benign augmentation
                remained disabled ({external_benign[0]['enabled']}) should also be stated clearly, because it keeps the
                final claim aligned with the dataset actually used.
                """
            )
        )
    )

    blocks.append(HeadingBlock(1, "14. Discussion: What the Journey Actually Demonstrated"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                Several conclusions follow from the chronology. First, data engineering was not a preparatory chore. It
                shaped the scientific result. The metadata merge enabled protocol-aware and family-aware analysis, and the
                later leakguard checks ensured that those metadata fields could guide evaluation without leaking into the
                feature set. Second, protocol heterogeneity was not a minor dataset characteristic. It changed which model
                families looked acceptable and ultimately justified routed inference as the correct operational design.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                Third, clean metrics were necessary but insufficient. The clean GPU and early HPO runs already showed that
                the classification task was easy to score well on. What separated the final result from the earlier ones
                was the shift from clean optimization to robust operating-envelope control. That shift only happened after
                benign-side WiFi attack exposed the real false-positive risk. Fourth, stability was not a luxury add-on. It
                directly changed the winner. If the project had selected by single-run rank, it would have promoted an
                XGBoost family that failed under rerun.
                """
            )
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The final methodological contribution is therefore broader than a single model artifact. The project built
                a reproducible selection process for a healthcare IoMT IDS under class imbalance, protocol heterogeneity,
                low-FPR constraints, and realistic adversarial pressure. That process is what makes the final model
                defensible. Without the journey, the ending would be a score. With the journey, it becomes a thesis result.
                """
            )
        )
    )

    script_rows = [
        [
            "merge_ciciomt_with_metadata.py",
            "Dataset construction",
            "Derives metadata, protocol hints, attack families, and deterministic profiling splits",
            "Makes later protocol-aware evaluation and leakage checks possible",
        ],
        [
            "generate_advanced_eda_report.py",
            "EDA",
            "Produces class balance, protocol composition, shift, and separability summaries",
            "Establishes why low-FPR and protocol slicing must drive the modeling agenda",
        ],
        [
            "train_baseline_models_stdlib.py",
            "Cheap learnability check",
            "Uses stable hash validation and reservoir sampling caps without third-party ML dependencies",
            "Proves the feature space is worth scaling before spending GPU budget",
        ],
        [
            "train_full_gpu_models.py",
            "First full-data GPU stage",
            "Selects threshold by maximizing recall under validation FPR <= 1%",
            "Shows that low-FPR intent on validation does not guarantee low test FPR",
        ],
        [
            "train_hpo_gpu_models.py",
            "Constrained HPO",
            "Optimizes validation F1 plus PR-AUC bonus with explicit FPR-gap penalty",
            "Improves optimization discipline but still leaves deployment credibility incomplete",
        ],
        [
            "train_hpo_gpu_models_fprfix.py",
            "Validation repair",
            "Builds protocol-and-class-aware file-level validation with controlled row fallback",
            "Reduces the chance that routed performance is flattered by split pathology",
        ],
        [
            "train_hpo_gpu_models_leakguard.py",
            "Leak prevention",
            "Checks metadata prefix, forbids leakage columns, and verifies train/test capture disjointness",
            "Turns strong routed metrics into trustworthy routed metrics",
        ],
        [
            "xgb_protocol_ids_utils.py",
            "Routed inference utility",
            "Loads per-protocol models and thresholds and exposes routed_predict",
            "Provides the operational bridge from per-protocol training to deployment and explainability",
        ],
        [
            "evaluate_xgb_robustness.py",
            "Robustness evaluation",
            "Applies realistic query attacks with semantic locks and relation-aware constraints",
            "Reveals benign-side WiFi false positives as the key weakness to harden",
        ],
        [
            "train_wifi_robust_hardening.py",
            "WiFi hardening",
            "Generates adversarial benign hard negatives and reselects threshold under constraints",
            "Transforms robustness from diagnosis into intervention",
        ],
        [
            "train_wifi_robust_rebalance_matrix.py",
            "WiFi gate selection",
            "Evaluates family variants under explicit clean-FPR, attacked-benign-FPR, and adv-recall gates",
            "Shows that gate passing, not score chasing, is the right robust selection rule",
        ],
        [
            "train_protocol_multimodel_robust_matrix.py",
            "Final robust search",
            "Warm-fits clean models, augments with hard negatives and adversarial rows, then selects thresholds lexicographically by gates",
            "Produces the stability-based final winner and the saved deployment artifact",
        ],
    ]

    blocks.append(HeadingBlock(1, "Appendix A. Script Verification Map"))
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                This appendix records how the chapter text was cross-checked against the actual codebase. The point is not
                to document every function in every file. The point is to show which scripts changed the scientific meaning
                of the results and what methodological behavior was verified directly in code.
                """
            )
        )
    )
    blocks.append(
        TableBlock(
            title="Table 12. Script-level verification used to write this chapter",
            columns=["Script", "Phase", "Verified behavior", "Why it mattered to the narrative"],
            rows=script_rows,
            note="The appendix table is intentionally selective. It keeps only the scripts that materially changed the argument of the thesis chapter.",
        )
    )
    blocks.append(
        ParagraphBlock(
            clean_text(
                """
                The chapter should therefore be read as a script-verified account of the experimentation process rather
                than as a retrospective summary written from memory. That is the main reason the final model claim is
                strong. The narrative has been made to follow the code and the saved artifacts, not the other way around.
                """
            )
        )
    )

    return blocks


def render_markdown(blocks: Sequence[Block]) -> str:
    parts: list[str] = []
    for block in blocks:
        if isinstance(block, TitleBlock):
            parts.append(f"# {block.title}\n\n{block.subtitle}")
        elif isinstance(block, HeadingBlock):
            parts.append(f"{'#' * block.level} {block.text}")
        elif isinstance(block, ParagraphBlock):
            parts.append(block.text)
        elif isinstance(block, TableBlock):
            table_text = f"**{block.title}**\n\n{markdown_table(block.columns, block.rows)}"
            if block.note:
                table_text += f"\n\n_Note: {block.note}_"
            parts.append(table_text)
        else:
            raise TypeError(f"Unsupported block type: {type(block)!r}")
    return "\n\n".join(parts) + "\n"


def set_run_font(run, font_name: str, size_pt: int, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, bold in [
        ("Title", 20, True),
        ("Subtitle", 12, False),
        ("Heading 1", 15, True),
        ("Heading 2", 13, True),
        ("Heading 3", 12, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold


def render_docx(blocks: Sequence[Block], out_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    first_title = True
    for block in blocks:
        if isinstance(block, TitleBlock):
            if not first_title:
                doc.add_section(WD_SECTION_START.NEW_PAGE)
            first_title = False
            title_par = doc.add_paragraph(style="Title")
            title_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_par.add_run(block.title)
            set_run_font(title_run, "Times New Roman", 20, bold=True)

            subtitle_par = doc.add_paragraph(style="Subtitle")
            subtitle_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_par.add_run(block.subtitle)
            set_run_font(subtitle_run, "Times New Roman", 12)
            doc.add_page_break()
        elif isinstance(block, HeadingBlock):
            level = min(max(block.level, 1), 3)
            par = doc.add_paragraph(style=f"Heading {level}")
            run = par.add_run(block.text)
            set_run_font(run, "Times New Roman", 16 - level, bold=True)
        elif isinstance(block, ParagraphBlock):
            par = doc.add_paragraph(style="Normal")
            run = par.add_run(block.text)
            set_run_font(run, "Times New Roman", 12)
        elif isinstance(block, TableBlock):
            caption = doc.add_paragraph(style="Normal")
            caption_format = caption.paragraph_format
            caption_format.space_before = Pt(8)
            caption_format.space_after = Pt(4)
            caption_run = caption.add_run(block.title)
            set_run_font(caption_run, "Times New Roman", 11, bold=True)

            table = doc.add_table(rows=1, cols=len(block.columns))
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            for idx, col in enumerate(block.columns):
                header_cells[idx].text = col
                for paragraph in header_cells[idx].paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, "Times New Roman", 10, bold=True)

            for row in block.rows:
                cells = table.add_row().cells
                for idx, cell_value in enumerate(row):
                    cells[idx].text = str(cell_value)
                    for paragraph in cells[idx].paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run, "Times New Roman", 10)

            if block.note:
                note_par = doc.add_paragraph(style="Normal")
                note_run = note_par.add_run(f"Note. {block.note}")
                set_run_font(note_run, "Times New Roman", 10, italic=True)
        else:
            raise TypeError(f"Unsupported block type: {type(block)!r}")

    doc.save(out_path)


def count_words(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))


def main() -> None:
    blocks = build_blocks()
    markdown = render_markdown(blocks)
    OUT_MD.write_text(markdown, encoding="utf-8")
    render_docx(blocks, OUT_DOCX)
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")
    print(f"Approximate words: {count_words(markdown):,}")


if __name__ == "__main__":
    main()
